"""Document Writer View - Generate resumes, cover letters, and cold emails"""

import flet as ft
from ui.styles.theme import AppTheme
from ui.components.file_uploader import FileUploadComponent
from services.document_service import DocumentService
from services.resume_service import ResumeService
from services.jd_service import JobDescriptionService
from core.auth import SessionManager

class WriterView:
    """View for AI-powered document generation"""
    
    def __init__(self, page: ft.Page):
        self.page = page
        self.user_id = SessionManager.get_user_id()
        self.current_document = None
        self.current_document_content = None
        self.current_document_type = None
        self.selected_resume_id = None
        self.resume_uploader = None
        self.save_file_picker = None
        self.resume_jd_dropdown = None
        self.cover_letter_jd_dropdown = None
        self.cold_email_jd_dropdown = None
        
    def _refresh_jd_dropdowns(self):
        """Refresh all JD dropdowns with latest JDs"""
        from services.jd_service import JobDescriptionService
        jds = JobDescriptionService.get_user_job_descriptions(self.user_id) or []
        jd_options = [
            ft.dropdown.Option(str(jd['jd_id']), 
                f"{jd.get('company_name', 'Company')} - {jd.get('job_title', 'Position')}") 
            for jd in jds
        ]
        
        if self.resume_jd_dropdown:
            self.resume_jd_dropdown.options = jd_options.copy()
            self.resume_jd_dropdown.disabled = len(jds) == 0
            self.resume_jd_dropdown.hint_text = "Choose a JD" if jds else "Add a job description first"
            self.resume_jd_dropdown.update()
            
        if self.cover_letter_jd_dropdown:
            self.cover_letter_jd_dropdown.options = jd_options.copy()
            self.cover_letter_jd_dropdown.disabled = len(jds) == 0
            self.cover_letter_jd_dropdown.hint_text = "Choose a JD" if jds else "Add a job description first"
            self.cover_letter_jd_dropdown.update()
            
        if self.cold_email_jd_dropdown:
            self.cold_email_jd_dropdown.options = (
                [ft.dropdown.Option("", "None")] + jd_options.copy()
            )
            self.cold_email_jd_dropdown.update()
        
    def build(self) -> ft.Container:
        """Build writer view"""
        # Header
        header = ft.Container(
            content=ft.Column([
                ft.Text("📝 Document Writer", size=28, weight=ft.FontWeight.BOLD),
                ft.Text("Generate optimized resumes, cover letters, and cold emails with AI",
                       size=14, color="grey")
            ], spacing=8),
            padding=AppTheme.PADDING_MEDIUM
        )
        
        # Resume selection section (shared across all tabs)
        self.resume_section = self._build_resume_section()
        
        # Tabs for different document types
        self.tabs = ft.Tabs(
            selected_index=0,
            tabs=[
                ft.Tab(text="Resume", icon=ft.Icons.DESCRIPTION),
                ft.Tab(text="Cover Letter", icon=ft.Icons.MAIL),
                ft.Tab(text="Cold Email", icon=ft.Icons.EMAIL),
            ],
            on_change=self._on_tab_change
        )
        
        # Resume tab content
        self.resume_tab = self._build_resume_tab()
        
        # Cover letter tab content
        self.cover_letter_tab = self._build_cover_letter_tab()
        
        # Cold email tab content
        self.cold_email_tab = self._build_cold_email_tab()
        
        # File picker for saving
        self.save_file_picker = ft.FilePicker(on_result=self._on_save_file_result)
        self.page.overlay.append(self.save_file_picker)
        
        # Document display area
        self.document_display = ft.Container(
            content=ft.Column([
                ft.Text("Generated Document", size=18, weight=ft.FontWeight.BOLD),
                ft.Divider(),
                ft.Text("Select options and click Generate to create a document", 
                       size=14, color="grey", italic=True)
            ], spacing=10),
            padding=20,
            border=ft.border.all(1, ft.Colors.OUTLINE),
            border_radius=8,
            bgcolor=ft.Colors.WHITE,
            expand=True
        )
        
        # Tab content container (shows one at a time)
        self.tab_content_container = ft.Container(
            content=self.resume_tab,
            width=400,
            padding=15
        )
        
        # Main content layout
        content = ft.Column([
            header,
            self.resume_section,
            ft.Divider(),
            self.tabs,
            ft.Row([
                self.tab_content_container,
                ft.VerticalDivider(width=1),
                self.document_display
            ], spacing=10, expand=True)
        ], spacing=10, expand=True)
        
        return ft.Container(
            content=content,
            padding=10,
            expand=True
        )
    
    def _build_resume_section(self) -> ft.Container:
        """Build resume upload/selection section"""
        # Get existing resumes
        resumes = ResumeService.get_all_resumes(self.user_id) or []
        
        # Resume dropdown
        self.resume_dropdown = ft.Dropdown(
            label="Select Resume",
            hint_text="Choose a resume" if resumes else "Upload a resume first",
            options=[ft.dropdown.Option(str(r['resume_id']), r.get('file_name', 'Resume')) for r in resumes],
            width=350,
            on_change=self._on_resume_selected,
            disabled=len(resumes) == 0
        )
        
        # Set default to active resume if available
        active_resume = ResumeService.get_active_resume(self.user_id)
        if active_resume and self.resume_dropdown.options:
            for opt in self.resume_dropdown.options:
                if opt.key == str(active_resume['resume_id']):
                    self.resume_dropdown.value = opt.key
                    self.selected_resume_id = active_resume['resume_id']
                    break
        
        # Resume uploader
        self.resume_uploader = FileUploadComponent(
            label="📤 Upload New Resume",
            allowed_extensions=['.pdf', '.docx', '.txt'],
            on_file_selected=self._on_resume_uploaded,
            help_text="PDF, DOCX, or TXT"
        )
        resume_uploader_ui = self.resume_uploader.build(page=self.page)
        
        # Ensure file picker is in overlay
        resume_file_picker = self.resume_uploader.get_file_picker()
        if resume_file_picker and resume_file_picker not in self.page.overlay:
            self.page.overlay.append(resume_file_picker)
        
        # Resume status
        self.resume_status_text = ft.Text(
            f"✓ {len(resumes)} resume(s) available" if resumes else "⚠ No resume found - upload a resume",
            size=11,
            color=ft.Colors.GREEN if resumes else ft.Colors.ORANGE,
            italic=True
        )
        
        return ft.Container(
            content=ft.Column([
                ft.Text("Resume Selection", size=14, weight=ft.FontWeight.BOLD),
                ft.Row([
                    self.resume_dropdown,
                    ft.Text("OR", size=12, color="grey", italic=True),
                    resume_uploader_ui
                ], spacing=10, wrap=True),
                self.resume_status_text
            ], spacing=8),
            padding=10,
            border=ft.border.all(1, ft.Colors.OUTLINE),
            border_radius=8,
            bgcolor=ft.Colors.GREY_50
        )
    
    def _on_resume_selected(self, e):
        """Handle resume selection from dropdown"""
        if e.control.value:
            self.selected_resume_id = int(e.control.value)
            self._update_resume_status()
            self.page.update()
    
    def _on_resume_uploaded(self, file_path: str, file_name: str):
        """Handle resume file upload"""
        try:
            # Upload resume
            resume_id = ResumeService.upload_resume(self.user_id, file_path)
            
            if resume_id:
                self.selected_resume_id = resume_id
                self._refresh_resume_dropdown()
                self._show_success(f"Resume '{file_name}' uploaded successfully!")
            else:
                self._show_error("Failed to upload resume")
        except Exception as ex:
            print(f"[ERROR] Error uploading resume: {ex}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Error: {str(ex)}")
    
    def _refresh_resume_dropdown(self):
        """Refresh resume dropdown with latest resumes"""
        resumes = ResumeService.get_all_resumes(self.user_id) or []
        
        self.resume_dropdown.options = [
            ft.dropdown.Option(str(r['resume_id']), r.get('file_name', 'Resume')) 
            for r in resumes
        ]
        self.resume_dropdown.disabled = len(resumes) == 0
        self.resume_dropdown.hint_text = "Choose a resume" if resumes else "Upload a resume first"
        
        # Select the newly uploaded resume
        if self.selected_resume_id:
            self.resume_dropdown.value = str(self.selected_resume_id)
        
        self._update_resume_status()
    
    def _update_resume_status(self):
        """Update resume status text"""
        resumes = ResumeService.get_all_resumes(self.user_id) or []
        if self.selected_resume_id:
            selected_resume = next((r for r in resumes if r['resume_id'] == self.selected_resume_id), None)
            resume_name = selected_resume.get('file_name', 'Resume') if selected_resume else 'Selected Resume'
            self.resume_status_text.value = f"✓ Using: {resume_name}"
            self.resume_status_text.color = ft.Colors.GREEN
        elif resumes:
            self.resume_status_text.value = f"✓ {len(resumes)} resume(s) available - please select one"
            self.resume_status_text.color = ft.Colors.ORANGE
        else:
            self.resume_status_text.value = "⚠ No resume found - upload a resume"
            self.resume_status_text.color = ft.Colors.ORANGE
    
    def _build_resume_tab(self) -> ft.Container:
        """Build resume generation tab"""
        # Get resumes and JDs
        resumes = ResumeService.get_all_resumes(self.user_id) or []
        jds = JobDescriptionService.get_user_job_descriptions(self.user_id) or []
        
        self.resume_jd_dropdown = ft.Dropdown(
            label="Select Job Description",
            hint_text="Choose a JD" if jds else "Add a job description first",
            options=[ft.dropdown.Option(str(jd['jd_id']), 
                    f"{jd.get('company_name', 'Company')} - {jd.get('job_title', 'Position')}") 
                    for jd in jds],
            width=350,
            disabled=len(jds) == 0
        )
        
        self.resume_template_dropdown = ft.Dropdown(
            label="Template Style",
            value="professional",
            options=[
                ft.dropdown.Option("professional", "Professional"),
                ft.dropdown.Option("modern", "Modern"),
                ft.dropdown.Option("creative", "Creative"),
                ft.dropdown.Option("ats", "ATS-Optimized"),
            ],
            width=350
        )
        
        self.generate_resume_btn = ft.ElevatedButton(
            "Generate Resume",
            icon=ft.Icons.AUTO_FIX_HIGH,
            on_click=self._on_generate_resume,
            disabled=len(jds) == 0,
            style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN, color=ft.Colors.WHITE)
        )
        
        return ft.Column([
            ft.Text("Resume Generator", size=16, weight=ft.FontWeight.BOLD),
            ft.Text("Generate an optimized resume tailored to a specific job", 
                   size=12, color="grey"),
            ft.Divider(height=10),
            self.resume_jd_dropdown,
            self.resume_template_dropdown,
            ft.Container(
                content=self.generate_resume_btn,
                padding=ft.padding.only(top=10)
            )
        ], spacing=10)
    
    def _build_cover_letter_tab(self) -> ft.Container:
        """Build cover letter generation tab"""
        jds = JobDescriptionService.get_user_job_descriptions(self.user_id) or []
        
        self.cover_letter_jd_dropdown = ft.Dropdown(
            label="Select Job Description",
            hint_text="Choose a JD" if jds else "Add a job description first",
            options=[ft.dropdown.Option(str(jd['jd_id']), 
                    f"{jd.get('company_name', 'Company')} - {jd.get('job_title', 'Position')}") 
                    for jd in jds],
            width=350,
            disabled=len(jds) == 0
        )
        
        self.cover_letter_length_dropdown = ft.Dropdown(
            label="Letter Length",
            value="medium",
            options=[
                ft.dropdown.Option("short", "Short (150-200 words)"),
                ft.dropdown.Option("medium", "Medium (250-350 words)"),
                ft.dropdown.Option("long", "Long (400-500 words)"),
            ],
            width=350
        )
        
        # Check if resume exists
        resume = ResumeService.get_active_resume(self.user_id)
        
        self.generate_cover_letter_btn = ft.ElevatedButton(
            "Generate Cover Letter",
            icon=ft.Icons.AUTO_FIX_HIGH,
            on_click=self._on_generate_cover_letter,
            disabled=len(jds) == 0,
            style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE, color=ft.Colors.WHITE)
        )
        
        return ft.Column([
            ft.Text("Cover Letter Generator", size=16, weight=ft.FontWeight.BOLD),
            ft.Text("Create a personalized cover letter based on your resume and job description", 
                   size=12, color="grey"),
            ft.Divider(height=10),
            self.cover_letter_jd_dropdown,
            self.cover_letter_length_dropdown,
            ft.Container(
                content=self.generate_cover_letter_btn,
                padding=ft.padding.only(top=10)
            )
        ], spacing=10)
    
    def _build_cold_email_tab(self) -> ft.Container:
        """Build cold email generation tab"""
        jds = JobDescriptionService.get_user_job_descriptions(self.user_id) or []
        
        self.cold_email_purpose_field = ft.TextField(
            label="Email Purpose",
            hint_text="e.g., networking, job inquiry, follow-up",
            value="networking",
            width=350
        )
        
        self.cold_email_company_field = ft.TextField(
            label="Company Name",
            hint_text="Enter company name",
            width=350
        )
        
        self.cold_email_recipient_dropdown = ft.Dropdown(
            label="Recipient Type",
            value="recruiter",
            options=[
                ft.dropdown.Option("recruiter", "Recruiter"),
                ft.dropdown.Option("hiring_manager", "Hiring Manager"),
                ft.dropdown.Option("employee", "Employee/Contact"),
                ft.dropdown.Option("executive", "Executive"),
            ],
            width=350
        )
        
        self.cold_email_jd_dropdown = ft.Dropdown(
            label="Related Job (Optional)",
            hint_text="Link to a job description",
            options=[ft.dropdown.Option("", "None")] + 
                    [ft.dropdown.Option(str(jd['jd_id']), 
                     f"{jd.get('company_name', 'Company')} - {jd.get('job_title', 'Position')}") 
                     for jd in jds],
            width=350
        )
        
        self.generate_cold_email_btn = ft.ElevatedButton(
            "Generate Cold Email",
            icon=ft.Icons.AUTO_FIX_HIGH,
            on_click=self._on_generate_cold_email,
            style=ft.ButtonStyle(bgcolor=ft.Colors.PURPLE, color=ft.Colors.WHITE)
        )
        
        return ft.Column([
            ft.Text("Cold Email Generator", size=16, weight=ft.FontWeight.BOLD),
            ft.Text("Create professional networking emails based on your resume", 
                   size=12, color="grey"),
            ft.Divider(height=10),
            self.cold_email_purpose_field,
            self.cold_email_company_field,
            self.cold_email_recipient_dropdown,
            self.cold_email_jd_dropdown,
            ft.Container(
                content=self.generate_cold_email_btn,
                padding=ft.padding.only(top=10)
            )
        ], spacing=10)
    
    def _on_tab_change(self, e):
        """Handle tab change"""
        tab_index = e.control.selected_index
        
        # Switch tab content
        if tab_index == 0:
            self.tab_content_container.content = self.resume_tab
        elif tab_index == 1:
            self.tab_content_container.content = self.cover_letter_tab
        elif tab_index == 2:
            self.tab_content_container.content = self.cold_email_tab
        
        # Clear document display when switching tabs
        self.document_display.content = ft.Column([
            ft.Text("Generated Document", size=18, weight=ft.FontWeight.BOLD),
            ft.Divider(),
            ft.Text("Select options and click Generate to create a document", 
                   size=14, color="grey", italic=True)
        ], spacing=10)
        
        self.page.update()
    
    def _on_generate_resume(self, e):
        """Generate resume"""
        if not self.resume_jd_dropdown.value:
            self._show_error("Please select a job description")
            return
        
        jd_id = int(self.resume_jd_dropdown.value)
        template = self.resume_template_dropdown.value
        
        # Show loading
        self.generate_resume_btn.disabled = True
        self.generate_resume_btn.text = "Generating..."
        self._show_loading("Generating optimized resume...")
        self.page.update()
        
        if not self.selected_resume_id:
            self._show_error("Please select or upload a resume first")
            return
        
        try:
            result = DocumentService.generate_resume(
                self.user_id,
                jd_id,
                template=template,
                resume_id=self.selected_resume_id
            )
            
            if result and "error" not in result:
                self._show_document(result.get('content', ''), "Resume")
                self._show_success("Resume generated successfully!")
            else:
                error_msg = result.get('error', 'Unknown error') if result else 'Generation failed'
                self._show_error(f"Error: {error_msg}")
        except Exception as ex:
            print(f"[ERROR] Error generating resume: {ex}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Error: {str(ex)}")
        finally:
            self.generate_resume_btn.disabled = False
            self.generate_resume_btn.text = "Generate Resume"
            self.page.update()
    
    def _on_generate_cover_letter(self, e):
        """Generate cover letter"""
        if not self.cover_letter_jd_dropdown.value:
            self._show_error("Please select a job description")
            return
        
        jd_id = int(self.cover_letter_jd_dropdown.value)
        length = self.cover_letter_length_dropdown.value
        
        # Show loading
        self.generate_cover_letter_btn.disabled = True
        self.generate_cover_letter_btn.text = "Generating..."
        self._show_loading("Generating cover letter...")
        self.page.update()
        
        if not self.selected_resume_id:
            self._show_error("Please select or upload a resume first")
            return
        
        try:
            result = DocumentService.generate_cover_letter(
                self.user_id,
                jd_id,
                length=length,
                resume_id=self.selected_resume_id
            )
            
            if result and "error" not in result:
                self._show_document(result.get('content', ''), "Cover Letter")
                self._show_success("Cover letter generated successfully!")
            else:
                error_msg = result.get('error', 'Unknown error') if result else 'Generation failed'
                self._show_error(f"Error: {error_msg}")
        except Exception as ex:
            print(f"[ERROR] Error generating cover letter: {ex}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Error: {str(ex)}")
        finally:
            self.generate_cover_letter_btn.disabled = False
            self.generate_cover_letter_btn.text = "Generate Cover Letter"
            self.page.update()
    
    def _on_generate_cold_email(self, e):
        """Generate cold email"""
        purpose = self.cold_email_purpose_field.value.strip()
        company = self.cold_email_company_field.value.strip()
        
        if not purpose:
            self._show_error("Please enter email purpose")
            return
        
        if not company:
            self._show_error("Please enter company name")
            return
        
        recipient_type = self.cold_email_recipient_dropdown.value
        jd_id = int(self.cold_email_jd_dropdown.value) if self.cold_email_jd_dropdown.value else None
        
        # Show loading
        self.generate_cold_email_btn.disabled = True
        self.generate_cold_email_btn.text = "Generating..."
        self._show_loading("Generating cold email...")
        self.page.update()
        
        if not self.selected_resume_id:
            self._show_error("Please select or upload a resume first")
            return
        
        try:
            result = DocumentService.generate_cold_email(
                self.user_id,
                purpose=purpose,
                company=company,
                recipient_type=recipient_type,
                jd_id=jd_id,
                resume_id=self.selected_resume_id
            )
            
            if result and "error" not in result:
                self._show_document(result.get('content', ''), "Cold Email")
                self._show_success("Cold email generated successfully!")
            else:
                error_msg = result.get('error', 'Unknown error') if result else 'Generation failed'
                self._show_error(f"Error: {error_msg}")
        except Exception as ex:
            print(f"[ERROR] Error generating cold email: {ex}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Error: {str(ex)}")
        finally:
            self.generate_cold_email_btn.disabled = False
            self.generate_cold_email_btn.text = "Generate Cold Email"
            self.page.update()
    
    def _show_loading(self, message: str):
        """Show loading state"""
        self.document_display.content = ft.Column([
            ft.ProgressRing(),
            ft.Text(message, size=14, color="grey")
        ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=10)
        self.document_display.update()
    
    def _show_document(self, content: str, doc_type: str):
        """Display generated document"""
        # Store document content and type
        self.current_document_content = content
        self.current_document_type = doc_type
        
        # Copy button
        copy_btn = ft.IconButton(
            icon=ft.Icons.COPY,
            tooltip="Copy to clipboard",
            on_click=lambda e: self._copy_to_clipboard(content)
        )
        
        # Export buttons
        export_txt_btn = ft.ElevatedButton(
            "Export TXT",
            icon=ft.Icons.DOWNLOAD,
            on_click=lambda e: self._export_document("txt"),
            style=ft.ButtonStyle(bgcolor=ft.Colors.BLUE_700, color=ft.Colors.WHITE)
        )
        
        export_docx_btn = ft.ElevatedButton(
            "Export DOCX",
            icon=ft.Icons.DOWNLOAD,
            on_click=lambda e: self._export_document("docx"),
            style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE)
        )
        
        # Document text area (editable)
        document_text = ft.TextField(
            value=content,
            multiline=True,
            min_lines=10,
            max_lines=30,
            expand=True,
            read_only=False,
            border_color=ft.Colors.OUTLINE,
            on_change=lambda e: self._on_document_edit(e.control.value)
        )
        
        self.document_display.content = ft.Column([
            ft.Row([
                ft.Text(f"{doc_type}", size=18, weight=ft.FontWeight.BOLD, expand=True),
                copy_btn
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            ft.Row([
                export_txt_btn,
                export_docx_btn
            ], spacing=10),
            ft.Divider(),
            document_text
        ], spacing=10, expand=True)
        self.document_display.update()
    
    def _on_document_edit(self, new_content: str):
        """Update stored content when user edits document"""
        self.current_document_content = new_content
    
    def _export_document(self, format_type: str):
        """Export document to file"""
        if not self.current_document_content:
            self._show_error("No document to export")
            return
        
        # Generate default filename
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_type_safe = self.current_document_type.lower().replace(" ", "_")
        
        if format_type == "txt":
            default_name = f"{doc_type_safe}_{timestamp}.txt"
            dialog_title = "Save as Text File"
        elif format_type == "docx":
            default_name = f"{doc_type_safe}_{timestamp}.docx"
            dialog_title = "Save as Word Document"
        else:
            self._show_error(f"Unsupported format: {format_type}")
            return
        
        # Store format for save callback
        self.pending_export_format = format_type
        
        # Use FilePicker to get save location
        # Flet 0.21+ supports save_file method
        try:
            # Check if save_file method exists
            if hasattr(self.save_file_picker, 'save_file'):
                self.save_file_picker.save_file(
                    dialog_title=dialog_title,
                    file_name=default_name,
                    allowed_extensions=[f".{format_type}"]
                )
            else:
                # Fallback: save to data/documents directory
                self._save_to_default_location(format_type, default_name)
        except Exception as e:
            print(f"[ERROR] Error opening save dialog: {e}")
            import traceback
            traceback.print_exc()
            # Fallback: save to default location
            self._save_to_default_location(format_type, default_name)
    
    def _on_save_file_result(self, e: ft.FilePickerResultEvent):
        """Handle file save result"""
        if not self.current_document_content:
            return
        
        format_type = getattr(self, 'pending_export_format', 'txt')
        
        # Get save path from event
        save_path = None
        if hasattr(e, 'path') and e.path:
            save_path = e.path
        elif hasattr(e, 'files') and e.files:
            # Some versions return files instead of path
            save_path = e.files[0].path if e.files else None
        
        if not save_path:
            # If no path provided, save to default location
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_type_safe = self.current_document_type.lower().replace(" ", "_")
            filename = f"{doc_type_safe}_{timestamp}.{format_type}"
            self._save_to_default_location(format_type, filename)
            return
        
        # Ensure correct extension
        if not save_path.endswith(f".{format_type}"):
            save_path = f"{save_path}.{format_type}"
        
        self._save_file_to_path(save_path, format_type)
    
    def _save_to_default_location(self, format_type: str, filename: str):
        """Save file to default documents directory"""
        import os
        from core.file_manager import FileManager
        
        # Ensure documents directory exists
        FileManager.ensure_directories()
        
        # Save to data/documents directory
        save_path = os.path.join(FileManager.DOCUMENTS_DIR, filename)
        self._save_file_to_path(save_path, format_type)
        
        # Show success with path
        import os
        abs_path = os.path.abspath(save_path)
        self._show_success(f"Document saved to: {abs_path}")
    
    def _save_file_to_path(self, file_path: str, format_type: str):
        """Save document content to file path"""
        try:
            if format_type == "txt":
                # Save as plain text
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(self.current_document_content)
                self._show_success(f"Document saved as TXT: {file_path}")
                
            elif format_type == "docx":
                # Save as DOCX using python-docx
                try:
                    from docx import Document
                    from docx.shared import Pt
                    
                    doc = Document()
                    
                    # Add title
                    title = doc.add_heading(self.current_document_type, 0)
                    
                    # Split content into paragraphs and add
                    paragraphs = self.current_document_content.split('\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            para = doc.add_paragraph(para_text.strip())
                            para.style.font.size = Pt(11)
                        else:
                            doc.add_paragraph()  # Empty line
                    
                    doc.save(file_path)
                    self._show_success(f"Document saved as DOCX: {file_path}")
                    
                except ImportError:
                    # Fallback to TXT if python-docx not available
                    file_path_txt = file_path.replace('.docx', '.txt')
                    with open(file_path_txt, 'w', encoding='utf-8') as f:
                        f.write(self.current_document_content)
                    self._show_error("DOCX export not available. Saved as TXT instead.")
                    
            else:
                self._show_error(f"Unsupported format: {format_type}")
                
        except Exception as e:
            print(f"[ERROR] Error saving file: {e}")
            import traceback
            traceback.print_exc()
            self._show_error(f"Error saving file: {str(e)}")
    
    def _copy_to_clipboard(self, text: str):
        """Copy text to clipboard"""
        self.page.set_clipboard(text)
        self._show_success("Copied to clipboard!")
    
    def _show_success(self, message: str):
        """Show success message"""
        self.page.snack_bar = ft.SnackBar(
            content=ft.Text(f"[OK] {message}"),
            bgcolor=ft.Colors.GREEN
        )
        self.page.snack_bar.open = True
        self.page.update()
    
    def _show_error(self, message: str):
        """Show error message"""
        self.page.snack_bar = ft.SnackBar(
            content=ft.Text(f"[ERROR] {message}"),
            bgcolor=ft.Colors.RED
        )
        self.page.snack_bar.open = True
        self.page.update()

